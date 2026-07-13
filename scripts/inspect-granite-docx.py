import argparse
import json
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("output")
    args = parser.parse_args()

    document = Document(args.input)
    payload = {
        "paragraphs": [
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": paragraph.text.strip(),
            }
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip()
        ],
        "tables": [
            {
                "index": table_index,
                "rows": [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                ],
            }
            for table_index, table in enumerate(document.tables)
        ],
    }

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"paragraphs={len(payload['paragraphs'])}")
    print(f"tables={len(payload['tables'])}")
    print(f"table_rows={sum(len(table['rows']) for table in payload['tables'])}")


if __name__ == "__main__":
    main()
